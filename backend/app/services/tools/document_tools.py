"""Doküman okuma tool'lari (FAZ 2.4): PDF, DOCX, Excel, Markdown, CSV, HTML."""
from __future__ import annotations

import csv
import io
import logging
from pathlib import Path
from typing import Any, Dict, List

from app.services.tools.base import BaseTool, ToolContext, ToolResult

logger = logging.getLogger(__name__)


def _read_pdf(path: Path, max_pages: int = 50) -> tuple[str, dict]:
    try:
        from pypdf import PdfReader  # type: ignore  # pyright: ignore[reportMissingImports]
    except ImportError:
        raise RuntimeError("pypdf kurulu degil: pip install pypdf")
    reader = PdfReader(str(path))
    pages = []
    total = len(reader.pages)
    for i, page in enumerate(reader.pages[:max_pages]):
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:  # pragma: no cover
            pages.append(f"(sayfa {i + 1} okunamadi: {exc})")
    text = "\n\n".join(pages)
    truncated = total > max_pages
    return text, {"format": "pdf", "page_count": total, "extracted_pages": min(total, max_pages), "truncated": truncated}


def _read_docx(path: Path) -> tuple[str, dict]:
    try:
        from docx import Document  # type: ignore  # pyright: ignore[reportMissingImports]
    except ImportError:
        raise RuntimeError("python-docx kurulu degil: pip install python-docx")
    doc = Document(str(path))
    parts: List[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    # tablolar
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    text = "\n".join(parts)
    return text, {"format": "docx", "paragraphs": len(doc.paragraphs), "tables": len(doc.tables)}


def _read_excel(path: Path, max_rows_per_sheet: int = 200) -> tuple[str, dict]:
    try:
        from openpyxl import load_workbook  # type: ignore  # pyright: ignore[reportMissingModuleSource]
    except ImportError:
        raise RuntimeError("openpyxl kurulu degil: pip install openpyxl")
    wb = load_workbook(str(path), read_only=True, data_only=True)
    parts: List[str] = []
    sheet_info = {}
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"\n## Sheet: {sheet_name}\n")
        rows_read = 0
        for row in ws.iter_rows(values_only=True):
            if rows_read >= max_rows_per_sheet:
                parts.append(f"...[{ws.max_row - max_rows_per_sheet} satir kesildi]")
                break
            cells = [str(c) if c is not None else "" for c in row]
            parts.append(" | ".join(cells))
            rows_read += 1
        sheet_info[sheet_name] = {"rows": rows_read, "max_row": ws.max_row}
    wb.close()
    text = "\n".join(parts)
    return text, {"format": "xlsx", "sheets": sheet_info}


def _read_csv(path: Path, max_rows: int = 500) -> tuple[str, dict]:
    parts: List[str] = []
    rows = 0
    with path.open("r", encoding="utf-8", errors="replace", newline="") as f:
        reader = csv.reader(f)
        for row in reader:
            if rows >= max_rows:
                parts.append(f"...[{rows} satir gosterildi, kalan kesildi]")
                break
            parts.append(" | ".join(row))
            rows += 1
    return "\n".join(parts), {"format": "csv", "rows": rows}


def _read_text(path: Path, max_chars: int = 16000) -> tuple[str, dict]:
    text = path.read_text(encoding="utf-8", errors="replace")
    fmt = path.suffix.lower().lstrip(".") or "txt"
    truncated = False
    if len(text) > max_chars:
        text = text[:max_chars] + f"\n...[{len(text) - max_chars} karakter kesildi]"
        truncated = True
    return text, {"format": fmt, "truncated": truncated}


def _read_html(path: Path, max_chars: int = 12000) -> tuple[str, dict]:
    raw = path.read_text(encoding="utf-8", errors="replace")
    try:
        from readability import Document  # type: ignore  # pyright: ignore[reportMissingImports]
        from lxml import html as lxml_html
        doc = Document(raw)
        tree = lxml_html.fromstring(doc.summary())
        text = tree.text_content().strip()
        title = doc.short_title() or ""
    except Exception:
        # fallback: tag stripping
        import re
        text = re.sub(r"<[^>]+>", " ", raw)
        text = re.sub(r"\s+", " ", text).strip()
        title = ""

    if len(text) > max_chars:
        text = text[:max_chars] + f"\n...[{len(text) - max_chars} karakter kesildi]"
    if title:
        text = f"# {title}\n\n{text}"
    return text, {"format": "html", "title": title}


class ReadDocumentTool(BaseTool):
    name = "read_document"
    description = (
        "Bir dokumani okur ve metnini cikarir. Desteklenen formatlar: "
        "PDF, DOCX, XLSX, XLS, CSV, MD, TXT, HTML, JSON, LOG."
    )
    permission = "file_system"
    parameters = {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "Okunacak dosyanin tam yolu"},
            "max_chars": {
                "type": "integer",
                "description": "Maksimum cikti uzunlugu (varsayilan 16000)",
                "default": 16000}},
        "required": ["path"]}

    async def execute(self, args: Dict[str, Any], context: ToolContext) -> ToolResult:
        path_str = (args.get("path") or "").strip()
        if not path_str:
            return ToolResult(ok=False, error="path gerekli")
        path = Path(path_str).expanduser().resolve()
        if not path.exists():
            return ToolResult(ok=False, error=f"Dosya yok: {path}")
        if not path.is_file():
            return ToolResult(ok=False, error=f"Dosya degil: {path}")

        max_chars = int(args.get("max_chars", 16000))
        ext = path.suffix.lower()

        try:
            if ext == ".pdf":
                text, meta = _read_pdf(path)
            elif ext == ".docx":
                text, meta = _read_docx(path)
            elif ext in (".xlsx", ".xlsm", ".xltx", ".xltm"):
                text, meta = _read_excel(path)
            elif ext == ".csv":
                text, meta = _read_csv(path)
            elif ext in (".html", ".htm"):
                text, meta = _read_html(path)
            elif ext in (".md", ".markdown", ".txt", ".log", ".json", ".yaml", ".yml", ".ini", ".toml", ".rst"):
                text, meta = _read_text(path)
            else:
                # bilinmeyenleri text olarak dene
                text, meta = _read_text(path)
        except RuntimeError as exc:
            return ToolResult(ok=False, error=str(exc))
        except Exception as exc:
            logger.exception("Dokuman okuma hatasi: %s", path)
            return ToolResult(ok=False, error=f"Okuma hatasi: {exc}")

        if len(text) > max_chars:
            text = text[:max_chars] + f"\n...[{len(text) - max_chars} karakter kesildi]"
            meta["truncated"] = True

        meta["path"] = str(path)
        meta["size_bytes"] = path.stat().st_size

        return ToolResult(
            ok=True,
            output=text or "(bos icerik)",
            data=meta,
        )